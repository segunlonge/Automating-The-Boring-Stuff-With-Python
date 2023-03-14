'''
Say you have a text file of guest names. This guests.txt file has one name per line, as follows:
Prof. Plum
Miss Scarlet
Col. Mustard
Al Sweigart
RoboCop

Write a program that would generate a Word document with custom invitations that look like Figure 13-11.
Since Python-Docx can use only those styles that already exist in the Word document,
you will have to first add these styles to a blank Word file and then open that file with Python-Docx.
There should be one invitation per page in the resulting Word document,
so call add_break() to add a page break after the last paragraph of each invitation.
This way, you will need to open only one Word document to print all of the invitations at once.
'''


import docx, os
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document

os.chdir('C:\Python Files\Automating The Boring Stuff With Python\Chapter 13 - Working with PDF and word Documents\Practice Projects - Working with PDF and Word Documents')
path = os.getcwd()

#doc = docx.Document(path+'\Invitation_Template.docx')
#print(doc.paragraphs[1].text)

guest_names = []

with open(path+'\guests.txt') as f:
    lines = [line.rstrip('\n') for line in f]
    guest_list = lines
    #print(guest_list)

master_switch = 0

for l in guest_list:
    #print(l)
    doc = docx.Document(path+'\Invitation_Template.docx')
    #print(doc.paragraphs[1].runs[0].font.size)
    #print(doc.paragraphs[1].runs[0].font.name)
    doc.paragraphs[1].text = l
    doc.paragraphs[1].runs[0].bold = True
    doc.paragraphs[1].runs[0].font.size = Pt(36)
    doc.save(l+'.docx')

    #print(master_switch)

    if master_switch == 0:
        master = Document(l+'.docx')
        master_file = l+'.docx'
        composer = Composer(master)
        master_switch = 1

        #print(master)

    if master_file != l+'.docx':
        print(master_file)
        print(l+'.docx')
        doc1 = Document(l+'.docx')
        composer.append(doc1)

composer.save("combined.docx")
    

